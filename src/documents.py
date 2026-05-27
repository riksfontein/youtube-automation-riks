"""
documents.py — Generate the 5 production Word documents per video.
Doc 1: Full scene brief (narration + image prompt + all video prompts)
Doc 2: Image prompts only, blank-line separated, NBP reference per scene
Doc 3A: Video prompts for Meta AI
Doc 3B: Video prompts for Google Flow VEO3.1
Doc 3C: Video prompts for Grok Aurora (6s or 10s)
"""

import os
import json
import requests
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]

CHANNEL_STYLES = {
    "AE": {
        "color_primary":   (14, 58, 92),    # deep navy
        "color_accent":    (30, 92, 42),    # forest green
        "channel_name":    "Ancient Earth Cinema",
        "niche":           "prehistoric natural history documentary",
        "visual_grade":    "Netflix/BBC Earth Prehistoric Planet — deep saturated color, volumetric light",
        "caption_font":    "Bebas Neue",
        "caption_case":    "ALL CAPS"
    },
    "GIA": {
        "color_primary":   (26, 46, 74),    # dark navy
        "color_accent":    (125, 90, 10),   # amber
        "channel_name":    "The Global Intel Analyst",
        "niche":           "geopolitical intelligence documentary",
        "visual_grade":    "Narcos/Vice News — teal-orange grade, presence anchor, real specific locations",
        "caption_font":    "Bebas Neue",
        "caption_case":    "ALL CAPS"
    },
    "BF": {
        "color_primary":   (44, 24, 16),    # deep brown
        "color_accent":    (125, 90, 10),   # amber gold
        "channel_name":    "The AI Bible Forensic",
        "niche":           "biblical archaeology documentary",
        "visual_grade":    "Warm amber and deep shadow — ancient parchment tone",
        "caption_font":    "Cinzel",
        "caption_case":    "Mixed-case"
    }
}


def _heading(doc, text, level=1, color=(10, 22, 40)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = RGBColor(*color)
    return p


def _body(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _code_block(doc, text):
    """Add text in a light grey shaded paragraph (for prompts)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Courier New"
    return p


# ─────────────────────────────────────────────────────────────
# Scene generation via Claude
# ─────────────────────────────────────────────────────────────

def generate_scenes(script_text: str, channel: str, scene_timings: list[dict]) -> list[dict]:
    """
    Use Claude to generate full scene data for all documents.
    Returns list of scene dicts with all prompts.
    """
    style = CHANNEL_STYLES[channel]

    # Build reference context from scene timings
    scene_timing_context = json.dumps([
        {
            "scene_number": s["scene_number"],
            "text": s["text"],
            "duration_s": s["duration_s"],
            "grok_duration": s["grok_duration"]
        }
        for s in scene_timings
    ], indent=2)

    prompt = f"""You are generating production documents for a faceless AI YouTube documentary channel.

CHANNEL: {style['channel_name']}
NICHE: {style['niche']}
VISUAL GRADE: {style['visual_grade']}

SCENE TIMINGS (use these exact scene texts and numbers):
{scene_timing_context}

For EACH scene above, generate:
1. image_prompt: Full NBP Nano Banana Pro image prompt
   - Start with: "Cinematic {style['niche']} reference image, {style['visual_grade']}, [lens type], 35mm film grain."
   - Include: specific environment, named light source, scale anchor, presence anchor
   - End with: "The viewer is [specific position] at [specific time]. [Physical sensory detail]. [Scale or wonder detail]."
   - NO readable text in image
   - Fully detailed — no character limit

2. nbp_composite_ref: Which composite reference to load (e.g. "ae-composite-apex-deep-ocean" or "none")

3. video_prompt_meta: For Meta AI
   - Motion verb first, action-oriented, 2-3 short sentences
   - NO timestamp format
   - Describe primary motion clearly

4. video_prompt_veo31: For Google Flow VEO3.1
   - Structure: [Camera movement] [Subject] [Action] [Environment] [Lighting] [Style/Audio]
   - Include timestamp format for most scenes: [00:00-00:05] first angle, [00:05-00:08] second angle
   - Single angle (no timestamp) for environment-only and closing scenes
   - Humans/creatures minimally animated — environment carries motion
   - Include audio direction at end

5. video_prompt_grok: For Grok Aurora
   - Structure: [Subject] [Action/Motion] [Camera Movement] [Visual Style] [Audio Direction]
   - Subject FIRST always
   - Duration: {"6" if scene["grok_duration"] == 6 else "use scene grok_duration"} seconds for scenes ≤10 words
   - Duration: 10 seconds for scenes with 11-13 words
   - No timestamp format
   - Audio direction embedded naturally at end

Return a JSON array. Each element:
{{
  "scene_number": <int>,
  "text": "<exact scene text>",
  "duration_s": <float>,
  "grok_duration": <6 or 10>,
  "nbp_composite_ref": "<ref_id or none>",
  "image_prompt": "<full NBP prompt>",
  "video_prompt_meta": "<Meta AI prompt>",
  "video_prompt_veo31": "<VEO3.1 prompt>",
  "video_prompt_grok": "<Grok Aurora prompt>"
}}

Return ONLY valid JSON array. No markdown code fences."""

    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        },
        json={
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 16000,
            "messages": [{"role": "user", "content": prompt}]
        },
        timeout=120
    )
    resp.raise_for_status()
    text = resp.json()["content"][0]["text"].strip()
    if "```" in text:
        text = text.split("```")[1].replace("json", "").strip().rstrip("```").strip()

    scenes = json.loads(text)
    print(f"[Documents] Generated prompts for {len(scenes)} scenes")
    return scenes


# ─────────────────────────────────────────────────────────────
# Document builders
# ─────────────────────────────────────────────────────────────

def build_doc1(scenes: list[dict], channel: str, video_title: str, output_dir: Path) -> str:
    """Doc 1 — Full scene brief: narration + image prompt + all 3 video prompts."""
    style = CHANNEL_STYLES[channel]
    doc = Document()

    # Title
    _heading(doc, f"Doc 1 — Full Scene Brief", 1, style["color_primary"])
    _heading(doc, f"{style['channel_name']} — {video_title}", 2, style["color_accent"])
    _body(doc, f"Total scenes: {len(scenes)} | Caption font: {style['caption_font']} {style['caption_case']}")
    doc.add_paragraph()

    for scene in scenes:
        n = scene["scene_number"]
        dur = scene.get("duration_s", 0)
        grok_dur = scene.get("grok_duration", 6)

        # Scene header
        _heading(doc, f"SCENE {n:03d}  ·  {dur:.1f}s narration  ·  Grok: {grok_dur}s", 2, style["color_primary"])

        # Narration
        _body(doc, "NARRATION:", 9)
        _body(doc, scene.get("text", ""), 11)

        # NBP Reference
        ref = scene.get("nbp_composite_ref", "none")
        if ref and ref != "none":
            _body(doc, f"NBP Reference: {ref}", 9)

        doc.add_paragraph()

        # Image prompt
        _body(doc, "IMAGE PROMPT (Doc 2 / NBP):", 9)
        _code_block(doc, scene.get("image_prompt", ""))

        doc.add_paragraph()

        # Video prompts
        _body(doc, "VIDEO PROMPT — Meta AI (Doc 3A):", 9)
        _code_block(doc, scene.get("video_prompt_meta", ""))

        _body(doc, "VIDEO PROMPT — VEO3.1 Google Flow (Doc 3B):", 9)
        _code_block(doc, scene.get("video_prompt_veo31", ""))

        _body(doc, f"VIDEO PROMPT — Grok Aurora {grok_dur}s (Doc 3C):", 9)
        _code_block(doc, scene.get("video_prompt_grok", ""))

        doc.add_paragraph("─" * 80)

    path = output_dir / "Doc1_Full_Scene_Brief.docx"
    doc.save(str(path))
    print(f"[Documents] Doc 1 saved: {path}")
    return str(path)


def build_doc2(scenes: list[dict], channel: str, video_title: str, output_dir: Path) -> str:
    """
    Doc 2 — Image prompts only.
    Plain prompts in order, blank line between each.
    NBP composite reference noted per scene.
    No headers, no scene numbers in the paste-ready section.
    """
    style = CHANNEL_STYLES[channel]
    doc = Document()

    _heading(doc, f"Doc 2 — Image Prompts", 1, style["color_primary"])
    _heading(doc, f"{style['channel_name']} — {video_title}", 2, style["color_accent"])
    _body(doc, f"Paste into VEO Automation Chrome Extension (Google Flow / Nano Banana Pro)")
    _body(doc, f"Save folder: {video_title.replace(' ', '_')}_images")
    _body(doc, f"Files will auto-download as: 001_image.jpg, 002_image.jpg, etc.")
    _body(doc, f"NBP composite reference is noted above each prompt — load before pasting.")
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    for scene in scenes:
        n     = scene["scene_number"]
        ref   = scene.get("nbp_composite_ref", "none")
        fname = f"{n:03d}_image.jpg"

        # Reference note (light grey, small)
        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(f"[{fname}] NBP ref: {ref}")
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(150, 150, 150)

        # Image prompt — paste-ready
        _code_block(doc, scene.get("image_prompt", ""))

        # Blank line between scenes (for extension parsing)
        doc.add_paragraph()

    path = output_dir / "Doc2_Image_Prompts.docx"
    doc.save(str(path))
    print(f"[Documents] Doc 2 saved: {path}")
    return str(path)


def build_doc3a(scenes: list[dict], channel: str, video_title: str, output_dir: Path) -> str:
    """Doc 3A — Video prompts for Meta AI. Blank line between each."""
    style = CHANNEL_STYLES[channel]
    doc = Document()

    _heading(doc, f"Doc 3A — Meta AI Video Prompts", 1, style["color_primary"])
    _heading(doc, f"{style['channel_name']} — {video_title}", 2, style["color_accent"])
    _body(doc, "Paste into Auto Meta extension. One blank line between scenes.")
    _body(doc, "Motion verb first. Action-oriented. 2-3 sentences max.")
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    for scene in scenes:
        n     = scene["scene_number"]
        fname = f"{n:03d}_video"

        # File reference
        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(f"[{fname}.mp4]")
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(150, 150, 150)

        # Prompt — paste-ready
        _code_block(doc, scene.get("video_prompt_meta", ""))
        doc.add_paragraph()

    path = output_dir / "Doc3A_Meta_AI_Video_Prompts.docx"
    doc.save(str(path))
    print(f"[Documents] Doc 3A saved: {path}")
    return str(path)


def build_doc3b(scenes: list[dict], channel: str, video_title: str, output_dir: Path) -> str:
    """Doc 3B — Video prompts for Google Flow VEO3.1. Blank line between each."""
    style = CHANNEL_STYLES[channel]
    doc = Document()

    _heading(doc, f"Doc 3B — Google Flow VEO3.1 Video Prompts", 1, style["color_primary"])
    _heading(doc, f"{style['channel_name']} — {video_title}", 2, style["color_accent"])
    _body(doc, "Paste into VEO Automation extension (Google Flow). One blank line between scenes.")
    _body(doc, "Camera movement first. Timestamp format [00:00-00:05]/[00:05-00:08]. Max 8s clips.")
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    for scene in scenes:
        n     = scene["scene_number"]
        dur   = scene.get("duration_s", 0)
        fname = f"{n:03d}_video"

        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(f"[{fname}.mp4]  {dur:.1f}s narration → use 8s VEO3.1 clip")
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(150, 150, 150)

        _code_block(doc, scene.get("video_prompt_veo31", ""))
        doc.add_paragraph()

    path = output_dir / "Doc3B_VEO31_Video_Prompts.docx"
    doc.save(str(path))
    print(f"[Documents] Doc 3B saved: {path}")
    return str(path)


def build_doc3c(scenes: list[dict], channel: str, video_title: str, output_dir: Path) -> str:
    """Doc 3C — Video prompts for Grok Aurora. 6s or 10s per scene."""
    style = CHANNEL_STYLES[channel]
    doc = Document()

    _heading(doc, f"Doc 3C — Grok Aurora Video Prompts", 1, style["color_primary"])
    _heading(doc, f"{style['channel_name']} — {video_title}", 2, style["color_accent"])
    _body(doc, "Paste into Grok extension. One blank line between scenes.")
    _body(doc, "Subject FIRST. Duration: 6s (≤10 words) or 10s (11-13 words). No timestamp format.")
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    for scene in scenes:
        n          = scene["scene_number"]
        dur        = scene.get("duration_s", 0)
        grok_dur   = scene.get("grok_duration", 6)
        word_count = scene.get("word_count") or len(scene.get("text", "").split())
        fname      = f"{n:03d}_video"

        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(
            f"[{fname}.mp4]  {word_count}w / {dur:.1f}s narration → set Grok to {grok_dur}s"
        )
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(150, 150, 150)

        _code_block(doc, scene.get("video_prompt_grok", ""))
        doc.add_paragraph()

    path = output_dir / "Doc3C_Grok_Aurora_Video_Prompts.docx"
    doc.save(str(path))
    print(f"[Documents] Doc 3C saved: {path}")
    return str(path)


def generate_metadata(
    script_text: str,
    ab_titles: list[str],
    channel: str,
    rotation_name: str,
    chapter_timings: list[dict]
) -> dict:
    """
    Generate full YouTube metadata:
    description, tags, hashtags, chapters.
    """
    from src.tts import format_chapters_for_description

    style = CHANNEL_STYLES[channel]

    # Format chapters
    chapters_text = format_chapters_for_description(chapter_timings)

    # Generate description via Claude
    prompt = f"""Write a YouTube video description for:
Channel: {style['channel_name']}
Rotation: {rotation_name}
Title: {ab_titles[0] if ab_titles else 'Documentary Video'}

Script excerpt (first 400 words):
{script_text[:1600]}

Rules:
- 2-3 paragraph description, 150-250 words total
- First sentence hooks the viewer
- Second paragraph expands on what they will learn/discover
- Third paragraph is a call to action (subscribe, comment)
- Professional documentary tone
- Include relevant keywords naturally

Also generate:
- 12-15 tags (comma separated, no #)
- 12-15 hashtags (with #, include channel hashtag)

Return JSON:
{{
  "description": "full description text",
  "tags": ["tag1", "tag2", ...],
  "hashtags": ["#Hashtag1", "#Hashtag2", ...]
}}"""

    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        },
        json={
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 1024,
            "messages": [{"role": "user", "content": prompt}]
        },
        timeout=30
    )
    resp.raise_for_status()
    text = resp.json()["content"][0]["text"].strip()
    if "```" in text:
        text = text.split("```")[1].replace("json", "").strip().rstrip("```")

    meta = json.loads(text)

    # Assemble full description with chapters and hashtags
    full_description = meta["description"]
    full_description += f"\n\nCHAPTERS:\n{chapters_text}"
    hashtags_line = " ".join(meta.get("hashtags", [])[:15])
    full_description += f"\n\n{hashtags_line}"

    return {
        "description": full_description,
        "tags":        meta.get("tags", []),
        "hashtags":    meta.get("hashtags", []),
        "chapters":    chapters_text
    }


def build_all_documents(
    script_text: str,
    channel: str,
    video_title: str,
    scene_timings: list[dict],
    output_dir: Path
) -> dict:
    """
    Generate all 5 documents for a video.
    Returns dict with paths to all documents.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n[Documents] Building all 5 documents for: {video_title}")

    # Generate scene data — base prompts first
    scenes = generate_scenes(script_text, channel, scene_timings)

    # Enhance video prompts using storyboard skill (Phase 2)
    # This ensures cinematically coherent shot variety and three-act arc
    print("[Documents] Enhancing video prompts with storyboard skill...")
    try:
        from src.storyboard import generate_scene_video_prompts, generate_storyboard_sheet_prompt
        scenes = generate_scene_video_prompts(scenes, channel, script_text, video_title)

        # Generate Phase 1 storyboard sheet prompt
        storyboard_sheet = generate_storyboard_sheet_prompt(
            script_text, channel, scene_timings, video_title
        )
        if storyboard_sheet:
            sheet_path = output_dir / "Doc0_Storyboard_Sheet_Prompt.txt"
            sheet_path.write_text(storyboard_sheet, encoding="utf-8")
            print(f"[Documents] Storyboard sheet prompt saved: {sheet_path}")
    except Exception as e:
        print(f"[Documents] Storyboard enhancement skipped: {e}")

    # Build all 5 documents
    doc1_path = build_doc1(scenes, channel, video_title, output_dir)
    doc2_path = build_doc2(scenes, channel, video_title, output_dir)
    doc3a_path = build_doc3a(scenes, channel, video_title, output_dir)
    doc3b_path = build_doc3b(scenes, channel, video_title, output_dir)
    doc3c_path = build_doc3c(scenes, channel, video_title, output_dir)

    print(f"[Documents] All 5 documents complete")

    return {
        "doc1": doc1_path,
        "doc2": doc2_path,
        "doc3a": doc3a_path,
        "doc3b": doc3b_path,
        "doc3c": doc3c_path,
        "scenes": scenes
    }
